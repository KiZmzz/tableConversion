"""PDF to Word conversion backend."""

import os
import re
import tempfile
from pathlib import Path

from converters.exceptions import ConversionCancelled


def _ensure_docx_output(output_path):
    output = Path(output_path)
    if output.suffix.lower() != ".docx":
        output = output.with_suffix(".docx")
    if output.parent and not output.parent.exists():
        output.parent.mkdir(parents=True, exist_ok=True)
    return str(output)


def _is_cjk(text):
    return bool(re.search(r"[\u3400-\u9fff]", text or ""))


def _join_line_items(items, median_height):
    parts = []
    last_x1 = None

    for item in items:
        text = item["text"].strip()
        if not text:
            continue

        if not parts:
            parts.append(text)
            last_x1 = item["x1"]
            continue

        gap = item["x0"] - (last_x1 or item["x0"])
        prev = parts[-1]
        if gap > max(median_height * 2.2, 28):
            sep = "\t"
        elif _is_cjk(prev[-1:]) or _is_cjk(text[:1]):
            sep = ""
        elif prev[-1:].isalnum() and text[:1].isalnum():
            sep = " "
        else:
            sep = ""

        parts.append(sep + text)
        last_x1 = item["x1"]

    return "".join(parts).strip()


def _line_text_and_tabs(items, median_height, content_width_pt, image_width_px):
    parts = []
    tabs = []
    last_x1 = None

    for item in items:
        text = item["text"].strip()
        if not text:
            continue

        if not parts:
            parts.append(text)
            last_x1 = item["x1"]
            continue

        gap = item["x0"] - (last_x1 or item["x0"])
        prev = parts[-1]
        if gap > max(median_height * 2.2, 28):
            parts.append("\t" + text)
            tabs.append(item["x0"] / image_width_px * content_width_pt)
        elif _is_cjk(prev[-1:]) or _is_cjk(text[:1]):
            parts.append(text)
        elif prev[-1:].isalnum() and text[:1].isalnum():
            parts.append(" " + text)
        else:
            parts.append(text)
        last_x1 = item["x1"]

    return "".join(parts).strip(), tabs


def _group_ocr_items_into_lines(items):
    if not items:
        return []

    heights = sorted(max(1.0, item["y1"] - item["y0"]) for item in items)
    median_height = heights[len(heights) // 2]
    y_threshold = max(8.0, median_height * 0.65)

    lines = []
    for item in sorted(items, key=lambda x: (x["yc"], x["x0"])):
        matched = None
        for line in lines:
            if abs(item["yc"] - line["yc"]) <= y_threshold:
                matched = line
                break
        if matched is None:
            lines.append({"yc": item["yc"], "items": [item]})
        else:
            matched["items"].append(item)
            matched["yc"] = sum(i["yc"] for i in matched["items"]) / len(matched["items"])

    result = []
    for line in sorted(lines, key=lambda x: x["yc"]):
        line_items = sorted(line["items"], key=lambda x: x["x0"])
        text = _join_line_items(line_items, median_height)
        if text:
            result.append({
                "text": text,
                "items": line_items,
                "x0": min(item["x0"] for item in line_items),
                "x1": max(item["x1"] for item in line_items),
                "y0": min(item["y0"] for item in line_items),
                "y1": max(item["y1"] for item in line_items),
                "height": max(item["y1"] - item["y0"] for item in line_items),
            })
    return result


def _set_east_asian_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", font_name)


def _add_text_line(document, text, font_name="微软雅黑", font_size=10.5):
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    _set_east_asian_font(run, font_name)
    run.font.size = Pt(font_size)


def _add_layout_line(document, line, page_layout, previous_bottom_pt=None, font_name="微软雅黑"):
    from docx.shared import Pt

    image_width = page_layout["image_width"]
    content_width = page_layout["content_width_pt"]
    dpi = page_layout["dpi"]
    top_margin = page_layout["top_margin_pt"]

    line_top_pt = line["y0"] * 72.0 / dpi
    line_bottom_pt = line["y1"] * 72.0 / dpi
    line_height_pt = max(6.0, line_bottom_pt - line_top_pt)
    desired_y_pt = max(0.0, line_top_pt - top_margin)
    if previous_bottom_pt is None:
        spacing_before = max(0.0, desired_y_pt)
    else:
        spacing_before = max(0.0, desired_y_pt - previous_bottom_pt)

    font_size = min(18.0, max(7.0, line_height_pt * 0.78))
    left_indent = min(max(0.0, line["x0"] / image_width * content_width), max(0.0, content_width - 20.0))
    text, tabs = _line_text_and_tabs(line["items"], line["height"], content_width, image_width)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(left_indent)
    paragraph.paragraph_format.space_before = Pt(spacing_before * 0.82)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(max(font_size * 1.12, line_height_pt))
    for tab_pt in tabs:
        if tab_pt > left_indent + 4:
            paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(tab_pt))

    run = paragraph.add_run(text)
    _set_east_asian_font(run, font_name)
    run.font.size = Pt(font_size)

    return max(desired_y_pt + line_height_pt, (previous_bottom_pt or 0.0) + font_size * 1.12 + spacing_before)


def _render_page_to_image(page, image_path, dpi):
    import fitz

    zoom = dpi / 72.0
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    pix.save(image_path)


def _ocr_image(ocr_engine, image_path, min_confidence=0.35):
    result = ocr_engine(image_path)
    if not result or not result.txts:
        return []

    items = []
    for box, text, score in zip(result.boxes, result.txts, result.scores):
        text = (text or "").strip()
        if not text or float(score) < min_confidence:
            continue

        xs = [float(p[0]) for p in box]
        ys = [float(p[1]) for p in box]
        x0, x1 = min(xs), max(xs)
        y0, y1 = min(ys), max(ys)
        items.append({
            "text": text,
            "score": float(score),
            "x0": x0,
            "x1": x1,
            "y0": y0,
            "y1": y1,
            "yc": (y0 + y1) / 2,
        })
    return items


def _convert_pdf_to_editable_word(
    pdf_path,
    output_path,
    dpi,
    log,
    progress_callback=None,
    cancel_check=None,
):
    import fitz
    from docx import Document
    from docx.shared import Pt
    from rapidocr import RapidOCR

    if cancel_check:
        cancel_check()
    log("\n🔎 使用 OCR 可编辑 Word 模式...")
    log(f"🎯 初始化 OCR 引擎 (RapidOCR)，DPI={dpi}...")
    if progress_callback:
        progress_callback(0.12)

    ocr = RapidOCR()
    document = Document()
    section = document.sections[0]
    margin_pt = 18

    with fitz.open(pdf_path) as pdf, tempfile.TemporaryDirectory(prefix="pdf_to_word_ocr_") as tmpdir:
        total_pages = len(pdf)
        for page_index, page in enumerate(pdf):
            if cancel_check:
                cancel_check()
            page_no = page_index + 1
            if page_index == 0:
                section.page_width = Pt(page.rect.width)
                section.page_height = Pt(page.rect.height)
                section.top_margin = Pt(margin_pt)
                section.bottom_margin = Pt(margin_pt)
                section.left_margin = Pt(margin_pt)
                section.right_margin = Pt(margin_pt)

            image_path = os.path.join(tmpdir, f"page_{page_index}.png")
            log(f"\n  --- 第 {page_no}/{total_pages} 页：渲染 + OCR ---")
            _render_page_to_image(page, image_path, dpi)
            if progress_callback:
                progress_callback(0.15 + 0.2 * (page_no / total_pages))

            if cancel_check:
                cancel_check()
            items = _ocr_image(ocr, image_path)
            lines = _group_ocr_items_into_lines(items)
            log(f"  OCR 文本块：{len(items)} 个，重建文本行：{len(lines)} 行")

            if not lines:
                _add_text_line(document, f"[第 {page_no} 页未识别到文本]")
            else:
                page_layout = {
                    "dpi": dpi,
                    "image_width": page.rect.width * dpi / 72.0,
                    "content_width_pt": max(36.0, page.rect.width - margin_pt * 2),
                    "top_margin_pt": margin_pt,
                }
                previous_bottom_pt = None
                for line in lines:
                    if cancel_check:
                        cancel_check()
                    previous_bottom_pt = _add_layout_line(document, line, page_layout, previous_bottom_pt)

            if page_index < total_pages - 1:
                document.add_page_break()

            if progress_callback:
                progress_callback(0.35 + 0.55 * (page_no / total_pages))

    if cancel_check:
        cancel_check()
    log("\n📝 写入可编辑 Word 文档...")
    document.save(output_path)


def convert_pdf_to_word(
    pdf_path,
    output_path,
    dpi=300,
    log_callback=None,
    progress_callback=None,
    cancel_check=None,
):
    """Convert PDF to editable DOCX text using OCR."""
    def log(msg):
        print(msg)
        if log_callback:
            log_callback(msg)

    pdf_path = str(pdf_path)
    output_path = str(output_path)

    log(f"\n{'=' * 55}")
    log("  PDF 转 Word 工具")
    log(f"{'=' * 55}")
    log(f"\n📄 输入：{pdf_path}")
    log(f"📝 输出：{output_path}")

    if not os.path.exists(pdf_path):
        log(f"\n❌ 找不到文件 '{pdf_path}'")
        raise FileNotFoundError(f"找不到文件 '{pdf_path}'")

    if cancel_check:
        cancel_check()
    output_path = _ensure_docx_output(output_path)

    if progress_callback:
        progress_callback(0.05)

    try:
        _convert_pdf_to_editable_word(
            pdf_path,
            output_path,
            dpi,
            log,
            progress_callback,
            cancel_check=cancel_check,
        )
    except ConversionCancelled:
        log("\n⏹ 已停止当前任务。")
        raise

    if progress_callback:
        progress_callback(1.0)

    log(f"\n{'=' * 55}")
    log(f"  ✅ 完成！输出：{os.path.abspath(output_path)}")
    log(f"{'=' * 55}\n")
